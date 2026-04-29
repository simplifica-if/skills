"""
DocxReader - Módulo de leitura de arquivos DOCX

Responsável por ler arquivos DOCX e extrair elementos estruturados
como parágrafos, tabelas e formatação.
"""

from pathlib import Path
from typing import Iterator, List, Optional, Union
from dataclasses import dataclass, field
from enum import Enum

try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.text.hyperlink import Hyperlink
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    raise ImportError(
        "python-docx é necessário. Instale com: pip install python-docx"
    )


class ElementType(Enum):
    """Tipos de elementos do documento."""
    PARAGRAPH = "paragraph"
    TABLE = "table"
    HEADING = "heading"
    IMAGE = "image"


@dataclass
class TextRun:
    """Representa um trecho de texto com formatação."""
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    hyperlink_url: Optional[str] = None


@dataclass
class ParagraphElement:
    """Representa um parágrafo do documento."""
    type: ElementType = ElementType.PARAGRAPH
    runs: List[TextRun] = field(default_factory=list)
    style_name: Optional[str] = None
    is_list_item: bool = False
    list_level: int = 0
    list_type: Optional[str] = None  # 'bullet' ou 'number'

    @property
    def text(self) -> str:
        """Retorna o texto completo do parágrafo."""
        return "".join(run.text for run in self.runs)

    @property
    def is_empty(self) -> bool:
        """Verifica se o parágrafo está vazio."""
        return not self.text.strip()


@dataclass
class TableCell:
    """Representa uma célula de tabela."""
    text: str
    row_span: int = 1
    col_span: int = 1
    is_header: bool = False
    is_merged_continuation: bool = False


@dataclass
class TableElement:
    """Representa uma tabela do documento."""
    type: ElementType = ElementType.TABLE
    rows: List[List[TableCell]] = field(default_factory=list)

    @property
    def num_rows(self) -> int:
        return len(self.rows)

    @property
    def num_cols(self) -> int:
        return max(len(row) for row in self.rows) if self.rows else 0


@dataclass
class HeadingElement:
    """Representa um heading do documento."""
    text: str
    level: int
    type: ElementType = ElementType.HEADING
    style_name: Optional[str] = None


@dataclass
class ImageElement:
    """Representa uma imagem do documento."""
    type: ElementType = ElementType.IMAGE
    rId: str = ""                    # Relationship ID (rId5)
    partname: Optional[str] = None   # Caminho no ZIP (word/media/image1.png)
    alt_text: str = ""               # Texto alternativo
    width_emu: Optional[int] = None  # Largura em EMUs
    height_emu: Optional[int] = None # Altura em EMUs
    is_drawing: bool = False         # True se for shapes/drawings vetoriais (requer renderização)


DocumentElement = Union[ParagraphElement, TableElement, HeadingElement, ImageElement]


class DocxReader:
    """
    Leitor de arquivos DOCX.

    Lê um arquivo DOCX e extrai seus elementos em formato estruturado.

    Uso:
        reader = DocxReader("documento.docx")
        for element in reader.elements():
            if element.type == ElementType.PARAGRAPH:
                print(element.text)
    """

    # Mapeamento de estilos de heading do Word
    HEADING_STYLES = {
        'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
        'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
        'Title': 1,
        # Estilos em português
        'Título 1': 1, 'Título 2': 2, 'Título 3': 3,
        'Título 4': 4, 'Título 5': 5, 'Título 6': 6,
        'Título': 1,
    }

    # Estilos de lista do Word
    LIST_STYLES = {
        'List Bullet', 'List Number', 'List Paragraph',
        'Lista com marcadores', 'Lista numerada', 'Parágrafo de Lista',
    }

    def __init__(self, file_path: Union[str, Path]):
        """
        Inicializa o leitor com um arquivo DOCX.

        Args:
            file_path: Caminho para o arquivo DOCX

        Raises:
            FileNotFoundError: Se o arquivo não existir
            ValueError: Se o arquivo não for um DOCX válido
        """
        self.file_path = Path(file_path)

        if not self.file_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {self.file_path}")

        if self.file_path.suffix.lower() != '.docx':
            raise ValueError(f"Arquivo deve ser .docx: {self.file_path}")

        try:
            self._document = Document(str(self.file_path))
        except Exception as e:
            raise ValueError(f"Erro ao abrir documento DOCX: {e}")

    @property
    def document(self) -> DocxDocument:
        """Retorna o documento python-docx."""
        return self._document

    def elements(self) -> Iterator[DocumentElement]:
        """
        Itera sobre todos os elementos do documento na ordem.

        Yields:
            DocumentElement: Parágrafos, tabelas, headings e imagens
        """
        # O python-docx não mantém a ordem exata de parágrafos e tabelas,
        # então precisamos iterar pelo XML diretamente
        body = self._document.element.body

        for child in body:
            tag = child.tag.split('}')[-1]  # Remove namespace

            if tag == 'p':
                # Verificar se o parágrafo contém imagem
                image = self._extract_image_from_paragraph(child)
                if image:
                    yield image

                # Processar parágrafo normalmente
                para = self._get_paragraph_by_element(child)
                if para:
                    element = self._parse_paragraph(para)
                    if element:
                        yield element

            elif tag == 'tbl':
                # É uma tabela
                table = self._get_table_by_element(child)
                if table:
                    yield self._parse_table(table)

            elif tag == 'sdt':
                # Suporte a Content Controls (Google Docs)
                yield from self._process_sdt_content(child)

    def _process_sdt_content(self, sdt_element) -> Iterator[DocumentElement]:
        """
        Processa conteúdo dentro de Content Controls (w:sdt).

        Content Controls são usados pelo Google Docs para envolver conteúdo.
        Esta função extrai parágrafos e tabelas de dentro dessas estruturas.

        Args:
            sdt_element: Elemento XML w:sdt

        Yields:
            DocumentElement: Parágrafos, tabelas e headings encontrados
        """
        sdt_content = sdt_element.find(qn('w:sdtContent'))
        if sdt_content is None:
            return

        for inner_child in sdt_content:
            inner_tag = inner_child.tag.split('}')[-1]

            if inner_tag == 'p':
                # Criar Paragraph diretamente do elemento XML
                # (elementos dentro de sdt não aparecem em self._document.paragraphs)
                para = Paragraph(inner_child, self._document._body)
                element = self._parse_paragraph(para)
                if element:
                    yield element

            elif inner_tag == 'tbl':
                # Usar extração direta do XML para tabelas em sdt
                # (mais confiável que python-docx para estruturas complexas com VMERGE)
                yield self._parse_table_from_xml(inner_child)

            elif inner_tag == 'sdt':
                # Content Controls podem estar aninhados
                yield from self._process_sdt_content(inner_child)

    def _get_paragraph_by_element(self, element) -> Optional[Paragraph]:
        """Encontra o objeto Paragraph correspondente a um elemento XML."""
        for para in self._document.paragraphs:
            if para._element is element:
                return para
        return None

    def _get_table_by_element(self, element) -> Optional[Table]:
        """Encontra o objeto Table correspondente a um elemento XML."""
        for table in self._document.tables:
            if table._element is element:
                return table
        return None

    def _parse_paragraph(self, para: Paragraph) -> Optional[DocumentElement]:
        """
        Converte um parágrafo python-docx para nosso formato.

        Args:
            para: Parágrafo do python-docx

        Returns:
            HeadingElement, ParagraphElement ou None (se vazio)
        """
        style_name = para.style.name if para.style else None

        # Verificar se é heading por estilo
        if style_name in self.HEADING_STYLES:
            text = para.text.strip()
            if text:
                return HeadingElement(
                    text=text,
                    level=self.HEADING_STYLES[style_name],
                    style_name=style_name
                )
            return None

        # Verificar se é heading por formatação (fallback)
        heading = self._detect_heading_by_format(para)
        if heading:
            return heading

        # É um parágrafo normal
        runs = self._extract_runs(para)

        # Detectar se é item de lista
        is_list, list_level, list_type = self._detect_list_item(para)

        element = ParagraphElement(
            runs=runs,
            style_name=style_name,
            is_list_item=is_list,
            list_level=list_level,
            list_type=list_type
        )

        return element if not element.is_empty else None

    def _extract_runs(self, para: Paragraph) -> List[TextRun]:
        """Extrai runs com formatação de um parágrafo, incluindo hyperlinks."""
        runs = []

        for item in para.iter_inner_content():
            if isinstance(item, Hyperlink):
                # Extrair URL (pode falhar se relationship estiver quebrada)
                url = None
                try:
                    url = item.url or None
                except KeyError:
                    url = None

                for run in item.runs:
                    if not run.text:
                        continue
                    runs.append(TextRun(
                        text=run.text,
                        bold=run.bold or False,
                        italic=run.italic or False,
                        underline=run.underline is not None and run.underline != False,
                        strike=run.font.strike or False,
                        hyperlink_url=url,
                    ))
            elif isinstance(item, Run):
                if not item.text:
                    continue
                runs.append(TextRun(
                    text=item.text,
                    bold=item.bold or False,
                    italic=item.italic or False,
                    underline=item.underline is not None and item.underline != False,
                    strike=item.font.strike or False,
                ))

        return runs

    def _get_all_runs(self, para: Paragraph) -> List[Run]:
        """Coleta todos os Run objects de um parágrafo, incluindo runs dentro de hyperlinks."""
        all_runs = []
        for item in para.iter_inner_content():
            if isinstance(item, Hyperlink):
                all_runs.extend(item.runs)
            elif isinstance(item, Run):
                all_runs.append(item)
        return all_runs

    def _detect_heading_by_format(self, para: Paragraph) -> Optional[HeadingElement]:
        """
        Detecta heading por formatação (negrito + tamanho) ou numeração.

        Padrões detectados:
        - Texto com numeração: "1.", "1.1", "1.1.1", etc.
        - Texto todo em negrito com tamanho maior
        """
        text = para.text.strip()
        if not text:
            return None

        # Detectar por numeração de seção (1., 1.1, 1.1.1, etc.)
        import re
        section_pattern = r'^(\d+(?:\.\d+)*\.?)\s*(.+)$'
        match = re.match(section_pattern, text)

        if match:
            section_num = match.group(1)
            # Nível baseado no número de pontos
            level = section_num.count('.') + 1
            if section_num.endswith('.'):
                level = section_num.rstrip('.').count('.') + 1

            # Verificar se está em negrito
            all_runs = self._get_all_runs(para)
            all_bold = all(
                run.bold for run in all_runs
                if run.text and run.text.strip()
            )

            if all_bold and level <= 6:
                return HeadingElement(
                    text=text,
                    level=min(level, 6),
                    style_name=None
                )

        # Detectar por formatação (texto curto, todo negrito, tamanho maior)
        if len(text) < 200:  # Headings são geralmente curtos
            all_runs = self._get_all_runs(para)
            all_bold = all(
                run.bold for run in all_runs
                if run.text and run.text.strip()
            )

            # Verificar tamanho da fonte
            font_size = None
            for run in all_runs:
                if run.font.size:
                    font_size = run.font.size.pt
                    break

            # Se todo negrito e fonte maior que 12pt, provavelmente é heading
            if all_bold and font_size and font_size >= 14:
                # Determinar nível pelo tamanho
                if font_size >= 18:
                    level = 1
                elif font_size >= 16:
                    level = 2
                else:
                    level = 3

                return HeadingElement(
                    text=text,
                    level=level,
                    style_name=None
                )

        return None

    def _detect_list_item(self, para: Paragraph) -> tuple:
        """
        Detecta se o parágrafo é um item de lista.

        Returns:
            Tuple[bool, int, Optional[str]]: (é_lista, nível, tipo)
        """
        style_name = para.style.name if para.style else ""

        # Detectar por estilo
        if any(ls in style_name for ls in self.LIST_STYLES):
            # Determinar tipo pelo estilo
            list_type = 'number' if 'Number' in style_name or 'numerada' in style_name else 'bullet'

            # Determinar nível pela indentação
            level = 0
            if para.paragraph_format.left_indent:
                # Aproximação: cada nível ~= 0.5"
                inches = para.paragraph_format.left_indent.inches
                level = int(inches / 0.5)

            return True, level, list_type

        # Detectar por XML (numPr)
        numPr = para._element.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0
            return True, level, 'bullet'  # Tipo padrão

        # Detectar por padrão no texto (a., b., •, -, etc.)
        text = para.text.strip()
        import re

        # Lista numerada: "1.", "a)", "(1)", etc.
        if re.match(r'^[\d]+[.)]\s', text) or re.match(r'^[a-z][.)]\s', text, re.I):
            return True, 0, 'number'

        # Lista com marcadores: "•", "-", "*"
        if text.startswith(('•', '-', '*', '–', '—')) and len(text) > 2:
            return True, 0, 'bullet'

        return False, 0, None

    def _parse_table(self, table: Table) -> TableElement:
        """
        Converte uma tabela python-docx para nosso formato.

        Args:
            table: Tabela do python-docx

        Returns:
            TableElement com células processadas
        """
        rows = []

        for row_idx, row in enumerate(table.rows):
            cells = []

            for cell_idx, cell in enumerate(row.cells):
                # Detectar células mescladas
                tc = cell._tc

                # Vertical merge (rowspan)
                v_merge = tc.find(qn('w:vMerge'))
                is_v_merged_start = v_merge is not None and v_merge.get(qn('w:val')) == 'restart'
                # Continuação se val é None ou 'continue' (não é 'restart')
                is_v_merged_cont = v_merge is not None and v_merge.get(qn('w:val')) != 'restart'

                # Horizontal merge (colspan) - grid span
                grid_span = tc.find(qn('w:gridSpan'))
                col_span = int(grid_span.get(qn('w:val'))) if grid_span is not None else 1

                # Calcular row_span é mais complexo - fazemos estimativa
                row_span = 1

                # Extrair texto da célula
                text = cell.text.strip()

                # Primeira linha geralmente é cabeçalho
                is_header = row_idx == 0

                cells.append(TableCell(
                    text=text,
                    row_span=row_span,
                    col_span=col_span,
                    is_header=is_header,
                    is_merged_continuation=is_v_merged_cont
                ))

            rows.append(cells)

        # Se nenhuma célula tem conteúdo, python-docx pode ter falhado
        # no grid mapping (comum em tabelas com VMERGE complexo)
        has_content = any(
            cell.text for row in rows for cell in row
        )
        if not has_content and hasattr(table, '_tbl'):
            return self._parse_table_from_xml(table._tbl)

        return TableElement(rows=rows)

    def _parse_table_from_xml(self, tbl_element) -> TableElement:
        """
        Extrai tabela diretamente do XML quando python-docx falha.

        Usado para tabelas complexas com células mescladas (VMERGE/COLSPAN)
        que causam falha no grid mapping do python-docx.

        Também lida com células dentro de Content Controls (w:sdt), comum em
        documentos exportados do Google Docs.

        Args:
            tbl_element: Elemento XML w:tbl

        Returns:
            TableElement com células extraídas
        """
        rows = []

        for tr in tbl_element.findall(qn('w:tr')):
            cells = []

            # Buscar células tanto diretamente quanto dentro de sdt
            # (Google Docs wraps cells in sdt elements)
            for tc in tr.findall('.//' + qn('w:tc')):
                # Extrair texto de todos os parágrafos da célula
                text_parts = []
                for t_elem in tc.findall('.//' + qn('w:t')):
                    if t_elem.text:
                        text_parts.append(t_elem.text)
                text = ''.join(text_parts)

                # Detectar merge vertical e horizontal
                tc_pr = tc.find(qn('w:tcPr'))
                is_merged_cont = False
                col_span = 1

                if tc_pr is not None:
                    # Vertical merge
                    v_merge = tc_pr.find(qn('w:vMerge'))
                    if v_merge is not None:
                        val = v_merge.get(qn('w:val'))
                        # Continuação se val é None ou 'continue' (não é 'restart')
                        is_merged_cont = val != 'restart'

                    # Horizontal merge (gridSpan)
                    grid_span = tc_pr.find(qn('w:gridSpan'))
                    if grid_span is not None:
                        col_span = int(grid_span.get(qn('w:val')) or 1)

                cells.append(TableCell(
                    text=text.strip(),
                    col_span=col_span,
                    is_header=(len(rows) == 0),
                    is_merged_continuation=is_merged_cont
                ))

            if cells:
                rows.append(cells)

        return TableElement(rows=rows)

    def get_metadata(self) -> dict:
        """
        Extrai metadados do documento.

        Returns:
            dict: Metadados como título, autor, etc.
        """
        core = self._document.core_properties

        return {
            'title': core.title,
            'author': core.author,
            'subject': core.subject,
            'keywords': core.keywords,
            'created': core.created.isoformat() if core.created else None,
            'modified': core.modified.isoformat() if core.modified else None,
            'last_modified_by': core.last_modified_by,
        }

    def _extract_image_from_paragraph(self, p_element) -> Optional[ImageElement]:
        """
        Extrai imagem de um elemento parágrafo XML.

        Detecta dois tipos de imagens:
        1. Imagens bitmap (w:drawing com a:blip) - podem ser extraídas do ZIP
        2. Shapes/desenhos vetoriais (w:pict) - requerem renderização

        Args:
            p_element: Elemento XML w:p

        Returns:
            ImageElement ou None se não houver imagem
        """
        # Namespaces para elementos de imagem
        ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

        # 1. Verificar w:drawing (imagens bitmap)
        drawing = p_element.find('.//' + qn('w:drawing'))
        if drawing is not None:
            return self._extract_image_from_drawing(drawing, ns_a, ns_r, ns_wp)

        # 2. Verificar w:pict (shapes/VML - desenhos vetoriais)
        pict = p_element.find('.//' + qn('w:pict'))
        if pict is not None:
            # Shapes requerem renderização, marcar como is_drawing=True
            return ImageElement(is_drawing=True, alt_text="Desenho vetorial (requer renderização)")

        return None

    def _extract_image_from_drawing(self, drawing, ns_a: str, ns_r: str, ns_wp: str) -> Optional[ImageElement]:
        """
        Extrai informações de imagem de um elemento w:drawing.

        Args:
            drawing: Elemento XML w:drawing
            ns_a: Namespace DrawingML
            ns_r: Namespace Relationships
            ns_wp: Namespace WordprocessingDrawing

        Returns:
            ImageElement ou None
        """
        # Buscar blip (referência à imagem)
        blip = drawing.find('.//{%s}blip' % ns_a)
        if blip is None:
            # Pode ser um shape sem imagem bitmap
            return ImageElement(is_drawing=True, alt_text="Shape sem bitmap")

        # Obter relationship ID
        rId = blip.get('{%s}embed' % ns_r)
        if not rId:
            return None

        # Resolver partname via relationships
        partname = self._resolve_image_relationship(rId)

        # Extrair dimensões (em EMUs - English Metric Units)
        width = None
        height = None
        extent = drawing.find('.//{%s}extent' % ns_wp)
        if extent is not None:
            width = int(extent.get('cx', 0)) if extent.get('cx') else None
            height = int(extent.get('cy', 0)) if extent.get('cy') else None

        # Extrair texto alternativo
        alt_text = ""
        doc_pr = drawing.find('.//{%s}docPr' % ns_wp)
        if doc_pr is not None:
            alt_text = doc_pr.get('descr', '') or doc_pr.get('title', '')

        return ImageElement(
            rId=rId,
            partname=partname,
            alt_text=alt_text,
            width_emu=width,
            height_emu=height,
            is_drawing=False
        )

    def _resolve_image_relationship(self, rId: str) -> Optional[str]:
        """
        Resolve relationship ID para caminho do arquivo de imagem no ZIP.

        Args:
            rId: Relationship ID (ex: "rId5")

        Returns:
            Caminho da imagem no ZIP (ex: "/word/media/image1.png") ou None
        """
        try:
            rel = self._document.part.rels[rId]
            # Retorna o partname (caminho no ZIP)
            return str(rel.target_part.partname)
        except (KeyError, AttributeError):
            return None
